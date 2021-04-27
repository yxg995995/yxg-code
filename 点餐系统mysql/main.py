import sys
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import QApplication,QMainWindow
from PyQt5.QtWidgets import *
from Ui_register import Ui_Register
from DataBase import DataBase
from Ui_NowOrBook import Ui_Type
from Ui_table import Ui_Table
from Ui_menu import Ui_Menu
from Ui_check import Ui_Check
from Ui_pay import Ui_Pay
from Ui_end import Ui_my_Order
from Ui_select import Ui_have_select
from Ui_admin_register import Ui_Admin_register
from Ui_Admin import Ui_Admin
import yagmail
import random
from PyQt5.QtSql import QSqlDatabase  , QSqlQueryModel , QSqlQuery
from Ui_help import Ui_Help
from Ui_adminlogin import Ui_admin_login
from Ui_userlogin import Ui_user_login
from docx import Document
from docx.shared import Inches
class user_Login(QDialog):
    def __init__(self):
        super(user_Login,self).__init__()
        self.login=Ui_user_login()
        self.login.setupUi(self)
        self.db=DataBase()
        self.login.userLogin.clicked.connect(self.ui_show)
        self.login.linkReg.clicked.connect(self.register_user)
        self.login.linkAdmin.clicked.connect(self.to_Admin)
        
    def ui_show(self):
        password=self.login.password.text()
        username=self.login.username.text()
        if self.db.Verify(username, password):
            self.close()
            self.ui=Type(self.login.username.text())
            self.ui.show()
        else:
            reply=QMessageBox.information(self,'错误','账号或者密码错误，请重新输入')
    
    def register_user(self):
        self.close()
        self.register=Register_window()
        self.register.show()  
    
    
    def to_Admin(self):
        self.close()
        self.admin=admin_Login()
        self.admin.show()


class admin_Login(QDialog):
    def __init__(self):
        super(admin_Login, self).__init__()
        self.login=Ui_admin_login()
        self.login=Ui_admin_login()
        self.login.setupUi(self)
        self.db=DataBase()
        self.login.adminLogin.clicked.connect(self.admin_show)
        self.login.adminR.clicked.connect(self.admin_r)
        self.login.linku.clicked.connect(self.to_user)
        self.login.cancel.clicked.connect(self.close)

    def admin_show(self):
        password=self.login.password.text()
        username=self.login.adminname.text()
        if self.db.Verify_admin(username, password):
            self.close()
            self.admin=Admin()
            self.admin.show()
        else:
            reply=QMessageBox.information(self,'错误','账号或者密码错误，请重新输入')

    def admin_r(self):
        self.close()
        self.register=Register_Admin()
        self.register.show()
    
    def to_user(self):
        self.close()
        self.user=user_Login()
        self.user.show()


          
class Register_window(QMainWindow):
    def __init__(self):
        super(Register_window, self).__init__()
        self.db=DataBase()
        self.register=Ui_Register()
        self.register.setupUi(self)
        self.register.register_ok.clicked.connect(self.userJoin)
        self.register.cancle.clicked.connect(self.to_login)
    
    def userJoin(self):
        username=self.register.userName.text()
        phone=self.register.phone.text()            
        password=self.register.password.text()
        cp=self.register.cp.text()
        if self.db.find_user(username):
            reply=QMessageBox.information(self,'错误','用户已存在')
            self.register.userName.clear()
        elif not phone.isdigit():
            reply=QMessageBox.information(self,'错误','电话号码必须为整数')
            self.register.phone.clear()
        elif len(phone)!=11 :
            reply=QMessageBox.information(self,'错误','电话号码必须为11位')
            self.register.phone.clear()
        elif self.db.find_phone(phone):
            reply=QMessageBox.information(self,'错误','电话号码已存在')
            self.register.phone.clear()
        else:
            if len(password)<8:
                self.register.password.clear()
                reply=QMessageBox.information(self,'错误','密码小于8位，请重新输入')
            elif password.isdigit() or password.isalpha():
                self.register.password.clear() 
                reply=QMessageBox.information(self,'错误','密码不能是纯数字或者字母')    
            elif cp!=password:
                self.register.cp.clear()
                reply=QMessageBox.information(self,'错误','两次输入的密码不一致')
            else:
                if self.db.find_phone(phone):
                    self.register.phone.clear()
                    reply=QMessageBox.information(self,'错误','号码已经注册')
                else:
                    reply=QMessageBox.information(self,'恭喜','注册成功')
                    self.db.insertUser(username,password,phone)
                    self.login=user_Login()
                    self.login.login.username.setText(username)
                    self.login.login.password.setText(password)
                    self.close()
                    self.login.show()

    def to_login(self):
        self.login=user_Login()
        self.close()
        self.login.show()

class Register_Admin(QMainWindow):
    def __init__(self):
        super(Register_Admin, self).__init__()
        self.db=DataBase()
        self.key='随便想一个吧'
        self.register=Ui_Admin_register()
        self.register.setupUi(self)
        self.register.register_2.clicked.connect(self.Join)
        self.register.cancle.clicked.connect(self.to_login)
        self.register.get.clicked.connect(self.send)
    
    def send(self):
        email=self.register.adminE.text()
        if email=='':
            reply=QMessageBox.information(self,'错误','邮箱为空')
        elif self.db.find_email(email):
            self.register.adminE.clear()
            reply=QMessageBox.information(self,'错误','邮箱已经注册')
        else:
            self.S='abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ123456789'
            self.code=''
            for i in range(6):
                self.code+=random.choice(self.S)
            message='您正在申请管理员注册，这是发送给您的验证码:'+self.code+'，请勿给他人观看'
            yag = yagmail.SMTP(
                    host='smtp.qq.com', user='2931148038@qq.com',
                    password='fsdkjqccelqfdcde', smtp_ssl=True
                ).send(email, '验证码', message)
            reply=QMessageBox.information(self,'提示','邮件发送成功')
    def Join(self):
        username=self.register.adminN.text()
        email=self.register.adminE.text()
        password=self.register.adminP.text()
        if self.db.find_admin(username):
            reply=QMessageBox.information(self,'错误','用户已存在')
            self.register.adminN.clear()
        else:
            if len(password)<8:
                reply=QMessageBox.information(self,'错误','密码小于8位，请重新输入')
                self.register.adminP.clear()
            elif password.isdigit() or password.isalpha():
                reply=QMessageBox.information(self,'错误','密码不能是纯数字或者字母')
                self.register.adminP.clear()
            else:
                if self.register.yanzhen.text()!=self.code:
                    reply=QMessageBox.information(self,'错误','验证码输入错误')
                    self.code=''
                    for i in range(6):
                        self.code+=random.choice(self.S)
                if self.register.adminK.text()!=self.key:
                    reply=QMessageBox.information(self,'错误','密钥输入错误')
                else:
                    reply=QMessageBox.information(self,'恭喜','注册成功')
                    self.db.insertAdmin(username,password,email)
                    self.login=admin_Login()
                    self.login.login.adminname.setText(username)
                    self.login.login.password.setText(password)
                    self.close()
                    self.login.show()

    def to_login(self):
        self.login=admin_Login()
        self.close()
        self.login.show()


class Admin(QDialog):
    def __init__(self):
        super(Admin,self).__init__()
        self.admin=Ui_Admin()
        self.DB=DataBase()
        self.admin.setupUi(self)
        self.admin.order_btn.clicked.connect(self.showO)
        self.admin.menu_btn.clicked.connect(self.showM)
        self.admin.table_btn.clicked.connect(self.showT)
        self.admin.user_btn.clicked.connect(self.showU)
        self.admin.back.clicked.connect(self.login)
        self.admin.add_btn.clicked.connect(self.add)
        self.admin.delete_btn.clicked.connect(self.delete)
        self.admin.one_empty.clicked.connect(self.set_one)
        self.admin.all_empty.clicked.connect(self.set_all)
        self.queryModel = None
        self.showM()
        self.admin.SUM.setSegmentStyle(QLCDNumber.Flat)
        SUM=self.DB.getSUM()
        self.admin.SUM.display(SUM)
    
    def add(self):
        if self.tablename=='menu':
            name,ok=QInputDialog.getText(self,'新增','输入菜名:')
            if ok and name not in self.DB.get_menuName():
                Class,ok=QInputDialog.getText(self,'新增','输入类别:')
                if ok:
                    Pic,ok=QInputDialog.getText(self,'新增','输入照片路径:')
                    if ok:
                        price,ok=QInputDialog.getInt(self,'新增','输入价格:')
                        if ok and price>0:
                            introduce,ok=QInputDialog.getText(self,'新增','输入简介:')
                            if ok:
                                self.DB.insert(self.tablename,[name,Class,Pic,price,introduce])
                                replay=QMessageBox.information(self,'提示','新增成功')
                                self.showM()
                        else:
                            replay=QMessageBox.information(self,'提示','价格必须大于0')
            else:
                replay=QMessageBox.information(self,'提示','菜名已存在')
        else :
            username,ok=QInputDialog.getText(self,'新增','输入用户名:')
            if ok and not self.DB.find_user(username):
                password,ok=QInputDialog.getText(self,'新增','输入密码:')
                if ok and not password.isdigit() and not password.isalpha() and len(password)>8:
                    phone,ok=QInputDialog.getText(self,'新增','输入电话号码:')
                    if ok and not self.DB.find_phone(phone) and phone.isdigit() and len(phone)==11:
                        self.DB.insertUser(username, password, phone)
                        replay=QMessageBox.information(self,'提示','新增成功')
                        self.showU()
                    else:
                        reply=QMessageBox.information(self,'错误','电话号码已经存在或者电话号码小于11位\
                                                                  或电话号码不为纯数字')
                        
                else:
                    reply=QMessageBox.information(self,'错误','密码必须大于等于8位，且不能为纯数字或字母')
            else:
                reply=QMessageBox.information(self,'错误','用户已经存在')
    def delete(self):
        if self.tablename=='menu':
            items=self.DB.get_menuName()
            item,ok=QInputDialog.getItem(self,'删除','菜名',items,0,False)
            if ok and item:
                self.DB.delete_menu(item)
                replay=QMessageBox.information(self,'提示','删除成功')
                self.showM()
        if self.tablename=='user':
            items=self.DB.get_user()
            item,ok=QInputDialog.getItem(self,'删除','人名',items,0,False)
            if ok and item:
                self.DB.delete_user(item)
                replay=QMessageBox.information(self,'提示','删除成功')
                self.showU()
                
    def set_one(self):
        items=self.DB.get_tableid()
        item,ok=QInputDialog.getItem(self,'选择餐桌编号','餐桌编号',items,0,False)
        if ok and item:
            self.DB.update_table(item)
            replay=QMessageBox.information(self,'提示','更新成功')
            self.showT()
            
    def set_all(self):
        self.DB.setTable_empty()
        replay=QMessageBox.information(self,'提示','更新成功')
        self.showT()
    
    def login(self):
        self.close()
        self.login=admin_Login()
        self.login.show()
    
    def showO(self):
        self.tablename='all_order'
        self.setTableView(self.tablename)
        self.admin.add_btn.setEnabled(False)
        self.admin.delete_btn.setEnabled(False)
        self.admin.one_empty.setEnabled(False)
        self.admin.all_empty.setEnabled(False)
        self.queryModel.setHeaderData(0, Qt.Horizontal, "用户名")
        self.queryModel.setHeaderData(1, Qt.Horizontal, "订餐时间")
        self.queryModel.setHeaderData(2, Qt.Horizontal, "菜名")
        self.queryModel.setHeaderData(3, Qt.Horizontal, "价格")
        self.queryModel.setHeaderData(4, Qt.Horizontal, "数量")
        self.queryModel.setHeaderData(5, Qt.Horizontal, "总计")
    
    def showM(self):
        self.tablename='menu'
        self.setTableView(self.tablename)
        self.admin.add_btn.setEnabled(True)
        self.admin.delete_btn.setEnabled(True)
        self.admin.one_empty.setEnabled(False)
        self.admin.all_empty.setEnabled(False)
        self.queryModel.setHeaderData(0, Qt.Horizontal, "菜品名称")
        self.queryModel.setHeaderData(1, Qt.Horizontal, "类别")
        self.queryModel.setHeaderData(2, Qt.Horizontal, "图片")
        self.queryModel.setHeaderData(3, Qt.Horizontal, "价格")
        self.queryModel.setHeaderData(4, Qt.Horizontal, "介绍")
    
    def showT(self):
        self.tablename='table'
        self.admin.add_btn.setEnabled(False)
        self.admin.delete_btn.setEnabled(False)
        self.admin.one_empty.setEnabled(True)
        self.admin.all_empty.setEnabled(True)
        self.setTableView(self.tablename)
        self.queryModel.setHeaderData(0, Qt.Horizontal, "餐桌编号")
        self.queryModel.setHeaderData(1, Qt.Horizontal, "状态")
    
    def showU(self):
        self.tablename='user'
        self.setTableView(self.tablename)
        self.admin.add_btn.setEnabled(True)
        self.admin.delete_btn.setEnabled(True)
        self.admin.one_empty.setEnabled(False)
        self.admin.all_empty.setEnabled(False)
        self.queryModel.setHeaderData(0, Qt.Horizontal, "用户姓名")
        self.queryModel.setHeaderData(1, Qt.Horizontal, "密码")
        self.queryModel.setHeaderData(2, Qt.Horizontal, "联系电话")
    
    def setTableView(self,tablename):
        self.db=QSqlDatabase.addDatabase('QMYSQL')
        self.db.setHostName('localhost')
        self.db.setDatabaseName('xinxi')
        self.db.setUserName('yxg')
        self.db.setPassword('yxg579521..')
        self.db.open()
        self.queryModel=QSqlQueryModel(self.admin)
        self.recordQuery(tablename)
        self.admin.tableView.setModel(self.queryModel)
        
    def recordQuery(self,tablename):
        q="select * from xinxi."+tablename
        szQuery = (q)
        self.queryModel.setQuery(szQuery)     
    


class Type(QMainWindow):
    def __init__(self,name):
        super(Type, self).__init__()
        self.name=name
        self.ui=Ui_Type()
        self.ui.setupUi(self)
        self.ui.now_btn.clicked.connect(self.Now)
        self.ui.book_btn.clicked.connect(self.Book)
        
        
    def Now(self):
        self.close()
        self.table=Table(self.name)
        self.table.show()
    
    def Book(self):
        items=('9:00-10:30','10:30-12:00','12:00-13:30',
               '13:30-15:00','17:00-18:30','18:30-20:00',
               '20:00-21:30','21:30-23:00')
        item,ok=QInputDialog.getItem(self,'预约时间','时间段',items,0,False)
        if ok and item:
            self.close()
            self.table=Table(self.name)
            self.table.show()
        
        
class Table(QMainWindow):
    def __init__(self,name):
        super(Table, self).__init__()
        self.name=name
        self.db=DataBase()
        self.table=Ui_Table()
        self.table.setupUi(self)
        self.empty()
        self.table.btn_1.clicked.connect(lambda:self.showMenu(1))
        self.table.btn_2.clicked.connect(lambda:self.showMenu(2))
        self.table.btn_3.clicked.connect(lambda:self.showMenu(3))
        self.table.btn_4.clicked.connect(lambda:self.showMenu(4))
        self.table.btn_5.clicked.connect(lambda:self.showMenu(5))
        self.table.btn_6.clicked.connect(lambda:self.showMenu(6))
        self.table.btn_7.clicked.connect(lambda:self.showMenu(7))
        self.table.btn_8.clicked.connect(lambda:self.showMenu(8))
        self.table.btn_9.clicked.connect(lambda:self.showMenu(9))
        self.table.btn_10.clicked.connect(lambda:self.showMenu(10))
        self.table.btn_11.clicked.connect(lambda:self.showMenu(11))
        self.table.btn_12.clicked.connect(lambda:self.showMenu(12))
        self.table.btn_13.clicked.connect(lambda:self.showMenu(13))
        self.table.btn_14.clicked.connect(lambda:self.showMenu(14))
        self.table.btn_15.clicked.connect(lambda:self.showMenu(15))
        self.table.btn_16.clicked.connect(lambda:self.showMenu(16))
        self.table.btn_17.clicked.connect(lambda:self.showMenu(17))
        self.table.btn_18.clicked.connect(lambda:self.showMenu(18))
        self.table.btn_19.clicked.connect(lambda:self.showMenu(19))
        self.table.btn_20.clicked.connect(lambda:self.showMenu(20))
        self.table.btn_21.clicked.connect(lambda:self.showMenu(21))
        self.table.btn_22.clicked.connect(lambda:self.showMenu(22))
        self.table.btn_23.clicked.connect(lambda:self.showMenu(23))
        self.table.btn_24.clicked.connect(lambda:self.showMenu(24))
        
    def empty(self):
        for i in range(1,25):
            if not self.db.is_empty(i):
                eval('self.table.btn_'+str(i)).setStyleSheet("border-radius: 40px;  border: 2px groove gray;\n"
"background-color: rgb(255, 0, 0);\n"
"")
                eval('self.table.btn_'+str(i)+'.setEnabled(False)')
    
    
    def showMenu(self,n):
        self.db.set_full(n)
        self.close()
        self.menu=Menu(self.name)
        self.menu.show()

class Menu(QMainWindow):
    def __init__(self,name,parameters=None):
        super(Menu, self).__init__()
        self.name=name
        self.menu=Ui_Menu()
        if parameters==None:
            self.menu.setupUi(self)
        else:
            self.menu.setupUi(self,R=parameters[0],S=parameters[1],Pizza=parameters[2],
                          P=parameters[3],D=parameters[4])
        self.menu.order_btn.clicked.connect(self.showSelect)
        self.menu.over_btn.clicked.connect(self.showorder)
    
    def showSelect(self):
        self.select=Select(self.name)
        self.select.show()
        
    def showorder(self):
        self.menu.showOrder()
        self.close()
        self.check=Check(self.name)
        self.check.show()

class Select(QDialog):
    def __init__(self,name,parent=None):
        super(Select, self).__init__(parent) 
        self.select=Ui_have_select()
        self.select.setupUi(self)
        self.db=DataBase()
        self.name=name
        self.select.pushButton.clicked.connect(self.close)

    
class Check(QMainWindow):
    def __init__(self,name):
        super(Check, self).__init__()
        self.name=name
        self.db=DataBase()
        self.parameters=self.db.get_parameters()
        self.check=Ui_Check()
        self.check.setupUi(self)
        self.result=self.db.get_quantityAndprice()
        self.check.to_pay.clicked.connect(self.showPay)
        self.check.back_menu.clicked.connect(self.backMenu)
        self.check.quantity.display(self.result[0])
        self.check.SUM.display(self.result[1])
        
    def showPay(self):
        self.close()
        self.pay=Pay(self.name,self.result[1])
        self.pay.show()

    def backMenu(self):
        self.close()
        self.menu=Menu(self.name,self.parameters)
        self.menu.show()
        
        

class Pay(QDialog):
    def __init__(self,name,price):
        super(Pay, self).__init__()
        self.name=name
        self.db=DataBase()
        self.pay=Ui_Pay()
        self.pay.setupUi(self)
        self.pay.paynumber.display(price)
        self.pay.pay.clicked.connect(self.to_pay)
        self.pay.cancel.clicked.connect(self.close)
    
    def to_pay(self):
        self.close()
        self.db.insert_allOrder(self.name)
        self.end=End(self.name)
        self.end.show()
        
class End(QMainWindow):
    def __init__(self,name,tablename=None):
        super(End, self).__init__()
        self.db=DataBase()
        self.name=name
        self.tablename=tablename
        self.end=Ui_my_Order()
        self.end.setupUi(self)
        self.end.backmenu.clicked.connect(self.Back)
        self.end.print.clicked.connect(self.Print)
    
    def Back(self):
        self.db.set_parameters()
        self.parameters=self.db.get_parameters()
        self.close()
        self.menu=Menu(self.name,self.parameters)
        self.menu.show()
    
    def Print(self):
        from ticket import Print_ticket
        msg=Print_ticket()
        self.document = Document()
        self.document.add_heading('结账小票', level=1)
        p =self.document.add_paragraph(msg)
        self.document.save('账单.docx')
        reply=QMessageBox.information(self,'提示','打印成功')   


class Main(QMainWindow):
    def __init__(self):
        super(Main, self).__init__()
        self.login=user_Login()
        self.login.show()
    

if __name__=="__main__":
    app=QApplication(sys.argv)
    win=Main()
    sys.exit(app.exec_())
