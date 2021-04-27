import sys
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import QApplication,QMainWindow
from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import *
from PyQt5.QtCore import QUrl
from Ui_select import Ui_select
from Draw import Draw
from analysis import *
from Ui_math import Ui_Math
from docx import Document
from docx.shared import Inches
import time
class Select(QWidget):
    def __init__(self):
        super(Select, self).__init__()
        self.ui=Ui_select()
        self.ui.setupUi(self)
        self.ui.Math.clicked.connect(self.showMath)

    def showMath(self):
        self.close()
        self.math=Math()
        self.math.show()

        
class Math(QWidget):
    def __init__(self):
        super(Math,self).__init__()
        self.ui=Ui_Math()
        self.ui.setupUi(self)
        self.draw=Draw()
        #分段间隔默认为30
        self.ui.distance.setText('30')
        #分档线默认为120，90，60
        self.ui.one_2.setText('120')
        self.ui.two_2.setText('90')
        self.ui.three_2.setText('60')
        
        self.one()
        self.two()
        self.three()
        self.four()
        
        self.ui.draw.clicked.connect(self.two)
        self.ui.draw_4.clicked.connect(self.four)
        self.ui.Export_btn.clicked.connect(self.Exportone)
        
    def Exportone(self):
        self.ui.Export_btn.setEnabled(False)
        start=time.time()
        self.draw.print_png()
        self.document = Document()
        self.document.add_heading('数学成绩分析报告', 0)
        self.document.add_heading('成绩概括分析', level=1)
        self.document.add_picture('分段面积.png', width=Inches(6))
        self.document.add_picture('富文本饼.png', width=Inches(6))
        p =self.document.add_paragraph(self.ui.report1.toPlainText())
        self.document.add_heading('分段成绩分析', level=1)
        self.document.add_picture('分段成绩直方图.png', width=Inches(6))
        self.document.add_picture('分段成绩饼图.png', width=Inches(6))
        p =self.document.add_paragraph(self.ui.report2.toPlainText())
        self.document.add_heading('四率一分分析', level=1)
        self.document.add_picture('平均分.png', width=Inches(6))
        self.document.add_picture('四率图.png', width=Inches(6))
        p =self.document.add_paragraph(self.ui.report3.toPlainText())
        self.document.add_heading('分档达标线分析', level=1)
        self.document.add_picture('一档图.png', width=Inches(7))
        self.document.add_picture('二档图.png', width=Inches(7))
        p =self.document.add_paragraph(self.ui.report4.toPlainText())
        self.document.save('数学分析.docx')
        end=time.time()
        replay=QMessageBox.information(self,'提示','输出文档成功，共耗时:'+str(end-start)+'秒')
        self.ui.Export_btn.setEnabled(True)
        
    def one(self):
        feature,x,y,M=self.draw.math_zongti()
        url1='D:/python/UI/学生成绩可视化/成绩概括面积图.html'
        self.ui.webEngineView11.load(QUrl(url1))
        url2='D:/python/UI/学生成绩可视化/成绩概括饼图.html'

        self.ui.webEngineView12.load(QUrl(url2))
        self.ui.Max.display(feature['max'])
        self.ui.Min.display(feature['min'])
        self.ui.Mid.display(feature['50%'])
        self.ui.Std.display(feature['std'])
        self.ui.Avg.display(feature['mean'])
        self.ui.Dif.display(feature['std']/feature['mean'])
        report='这次考试最高分为：'+str(feature['max'])+'分，最低分为：'+str(feature['min'])\
                +'分,平均分为：'+str(feature['mean'])+'\n'+\
                '其中'+x[y.index(max(y))]+'的人数最多'+'一共有' +str(max(y))+'人'+'占比为：'+str(M*100)+'%'
        self.ui.report1.setPlainText(report)
        
        
    def two(self):
        distance=self.ui.distance.text()
        Max,Min,M,N=self.draw.math_FD(distance)
        self.ui.Max_2.setText(str(Max))
        self.ui.Min_2.setText(str(Min))
        url1='D:/python/UI/学生成绩可视化/分段成绩直方图.png'
        self.ui.webEngineView21.load(QUrl(url1))
        url2='D:/python/UI/学生成绩可视化/分段成绩饼图.png'
        self.ui.webEngineView22.load(QUrl(url2))
        url3='D:/python/UI/学生成绩可视化/表格图.html'
        self.ui.webEngineView23.load(QUrl(url3))
        report='人数最多的区间段为：'+M\
        +'\n人数最少的区间段位：'+N
        self.ui.report2.setPlainText(report)
        
    def three(self):
        p1,p2,p3,p4=self.draw.math_four()
        self.ui.excellent.display(p1)
        self.ui.fine.display(p2)
        self.ui.Pass.display(p3)
        self.ui.low.display(p4)        
        url1='D:/python/UI/学生成绩可视化/平均分图.html'
        self.ui.webEngineView31.load(QUrl(url1))
        url2='D:/python/UI/学生成绩可视化/四率图.html'
        self.ui.webEngineView32.load(QUrl(url2))
        report='优秀率为：'+str(p1)+'%'+',良好率为：'+str(p2)+'%\n'+\
        ',及格率为：'+str(p3)+',低分率为：'+str(p4)
        self.ui.report3.setPlainText(report)        
        
    def four(self):
        one=self.ui.one_2.text()
        two=self.ui.two_2.text()
        three=self.ui.three_2.text()
        r1,r2=self.draw.math_fendang(one, two, three)
        url1='D:/python/UI/学生成绩可视化/分档表格图.html'
        self.ui.webEngineView41.load(QUrl(url1))
        url2='D:/python/UI/学生成绩可视化/一档直方图.html'
        self.ui.webEngineView42.load(QUrl(url2))
        url3='D:/python/UI/学生成绩可视化/二档直方图.html'
        self.ui.webEngineView43.load(QUrl(url3))
        report='学生分档线分析，可反映本次考试学生分档达线情况,'+\
                '分档线默认为一档线120，二档线90，三档线60,如需要更改'+\
                '可以在程序内更改\n该次考试达线情况如上述表格所示，其中：\n'+\
                '本次考试一档线，二档线累计上线人数分别为'+str(r2[0])+'人,'+str(r2[2])+'人,'+\
                '累计上线率分别为'+str(r2[1])+'%,'+str(r2[-1])+'%,一档累积上线率优秀的班级为：'+\
                r1[0]+','+r1[1]+','+r1[2]+'一档累积上线率较差的班级为：'+r1[3]+','+r1[4]+','+r1[5]+\
                '二档累积上线率优秀的班级为：'+r1[6]+','+r1[7]+','+r1[8]+','+\
                '二档累积上线率较差的班级为：'+r1[9]+','+r1[10]+','+r1[11]
        self.ui.report4.setPlainText(report)
        


class Main(QWidget):
    def __init__(self):
        super(Main, self).__init__()
        self.ui=Select()
        self.ui.show()
        
    

if __name__=="__main__":
    app=QApplication(sys.argv)
    win=Main()
    sys.exit(app.exec_())